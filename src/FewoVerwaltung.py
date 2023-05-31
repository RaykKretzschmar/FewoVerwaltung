import tkinter as tk
from tkinter import filedialog
import csv
import datetime
from docx import Document

class Textersetzung():
    def replace_text(replacements, save_path, template_path="Rechnungsvorlage.docx"):
        doc = Document(template_path)

        # Go through each paragraph in the document
        for p in doc.paragraphs:
            for run in p.runs:
                for old_text, new_text in replacements.items():
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

        # Go through each table in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for old_text, new_text in replacements.items():
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)

        
        # Save the document
        doc.save(save_path)
    

class NeuerKundeDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)

        self.title("Neuer Kunde")

        self.labels = {
            "Anrede": tk.StringVar(),
            "Vorname": tk.StringVar(),
            "Nachname": tk.StringVar(),
            "Stadt": tk.StringVar(),
            "Postleitzahl": tk.StringVar(),
            "Straße": tk.StringVar(),
            "Hausnummer": tk.StringVar(),
            "Kundennummer": tk.StringVar(value=self.generate_kundennummer()),
        }

        for i, (k, v) in enumerate(self.labels.items()):
            tk.Label(self, text=f"{k} : ", anchor="e").grid(row=i, column=0, sticky="e")
            tk.Entry(self, textvariable=v, width=30).grid(row=i, column=1, sticky="ew")

        tk.Button(self, text="Speichern", command=self.save).grid(row=i+1, column=0)
        tk.Button(self, text="Abbrechen", command=self.destroy).grid(row=i+1, column=1)

        self.grid_columnconfigure(1, weight=1)  # make the second column expandable

    def save(self):
        with open('Kunden.csv', 'a', newline='') as f:
            writer = csv.writer(f)
            writer.writerow([v.get() for v in self.labels.values()])
        self.destroy()

    def generate_kundennummer(self):
        return datetime.datetime.now().strftime('%Y%m%d%H%M%S')


class AuswahlDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)

        self.title("Kunden")
        self.configure()

        self.labels = ["Anrede", "Vorname", "Nachname", "Stadt", "PLZ", "Straße", "Hausnummer", "Kundennummer"]
        self.checkboxes = []

        for i, label in enumerate(self.labels):
            label_widget = tk.Label(self, text=label, font=('Arial', 14))
            label_widget.grid(row=0, column=i, sticky='ew')
            self.grid_columnconfigure(i, weight=1)

        self.selected_data = None  # Store the selected data

        with open('Kunden.csv', 'r') as f:
            reader = csv.reader(f)
            for i, line in enumerate(reader, start=1):
                for j, element in enumerate(line):
                    data_widget = tk.Label(self, text=element, font=('Arial', 12))
                    data_widget.grid(row=i, column=j, sticky='ew')
                
                button = tk.Button(self, text="Rechnung", command=lambda line=line: self.select_data(line))
                button.grid(row=i, column=len(self.labels), sticky='ew')

            self.grid_rowconfigure(i, weight=1)

        self.abbruch_button = tk.Button(self, text="Abbrechen", command=self.destroy, font=('Arial', 12))
        self.abbruch_button.grid(row=i+1, column=0, sticky='ew')

    def select_data(self, data):
        self.selected_data = {self.labels[i]: val for i, val in enumerate(data)}
        EingabeDialog(self, self.selected_data)


class EingabeDialog(tk.Toplevel):
    def __init__(self, parent, data):
        super().__init__(parent)
        self.parent = parent

        self.title("Rechnung erstellen")

        # data should be a dictionary
        for i, (k, v) in enumerate(data.items()):
            tk.Label(self, text=f"{k} :", anchor="e").grid(row=i, column=0, sticky="e")
            tk.Label(self, text=v, anchor="w").grid(row=i, column=1, sticky="w")
            self.grid_rowconfigure(i, weight=1)  # make the row expandable

        self.new_labels = {
            "Datum": tk.StringVar(),
            "Rechnungsnummer": tk.StringVar(),
            "Anreisedatum": tk.StringVar(),
            "Abreisedatum": tk.StringVar(),
            "Name der Ferienwohnung": tk.StringVar(),
            "Preis pro Nacht": tk.StringVar(),
        }

        for j, (k, v) in enumerate(self.new_labels.items(), start=i+1):
            tk.Label(self, text=f"{k} : ", anchor="e").grid(row=j, column=0, sticky="e")
            tk.Entry(self, textvariable=v, width=30).grid(row=j, column=1, sticky="ew")
            self.grid_rowconfigure(j, weight=1)  # make the row expandable

        tk.Button(self, text="Rechnung erstellen", command=self.invoice).grid(row=j+1, column=1, sticky="ew")
        tk.Button(self, text="Abbrechen", command=self.destroy).grid(row=j+1, column=0, sticky="ew")

        self.grid_columnconfigure(0, weight=1)  # make the first column expandable
        self.grid_columnconfigure(1, weight=1)  # make the second column expandable

    def invoice(self):
        # Format for date input
        date_format = "%d.%m.%Y"
        tax_percent = 7

        # Convert the dates to datetime objects
        anreisedatum = datetime.datetime.strptime(self.new_labels['Anreisedatum'].get(), date_format)
        abreisedatum = datetime.datetime.strptime(self.new_labels['Abreisedatum'].get(), date_format)

        number_of_nights = (abreisedatum - anreisedatum).days
        price_per_night = float(self.new_labels['Preis pro Nacht'].get().replace(',', '.'))
        price_for_all_nights = number_of_nights*price_per_night
        total_price = price_for_all_nights
        tax_amount = tax_percent*total_price/100
        net_amount = total_price-tax_amount

        # Merge selected data and new data
        all_data = {**self.parent.selected_data, 
                    **{k: v.get() for k, v in self.new_labels.items()}, 
                    'AnzahlDerÜbernachtungen': str(number_of_nights), 
                    'PpN': f'{price_per_night:.2f}'.replace('.', ','),
                    'NdFeWo': self.new_labels['Name der Ferienwohnung'].get(),
                    'ÜNKosten': f'{price_for_all_nights:.2f}'.replace('.', ','),
                    'GesamtBetrag': f'{total_price:.2f}'.replace('.', ','),
                    'MwstBetrag': f'{tax_amount:.2f}'.replace('.', ','),
                    'NettoBetrag': f'{net_amount:.2f}'.replace('.', ','),
                    }

        # Open a dialog to select the invoice template
        template_file_path = filedialog.askopenfilename(
            title="Wählen Sie eine Rechnungsvorlage aus", 
            filetypes=(("Word-Dokumente", "*.docx"), ("Alle Dateien", "*.*"))
        )

        # Open a dialog to select where to save the invoice
        invoice_file_path = filedialog.asksaveasfilename(
            title="Wählen Sie, wo die Rechnung gespeichert werden soll",
            filetypes=(("Word-Dokumente", "*.docx"), ("Alle Dateien", "*.*")),
            defaultextension=".docx"
        )

        # Replace text in .docx
        Textersetzung.replace_text(all_data, save_path=invoice_file_path, template_path=template_file_path)

        self.destroy()


class VerwaltungListener(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("Fewo Kundenverwaltung by Rayk Kretzschmar")
        self.geometry("800x600")
        self.configure()

        # Add a headline
        self.headline = tk.Label(self, text="FeWo-Verwaltung von Rayk Kretzschmar", font=('Arial', 20))
        self.headline.grid(row=0, column=0, columnspan=2, padx=10, pady=10)

        # Adjust font, colors, and padding of the buttons
        self.neuerKundeButton = tk.Button(self, text="Neuer Kunde", command=self.open_neuerKundeDialog, font=('Arial', 14), bg='skyblue', fg='black', padx=5, pady=5)
        self.rechnungButton = tk.Button(self, text="Rechnung erstellen", command=self.open_auswahlDialog, font=('Arial', 14), bg='skyblue', fg='black', padx=5, pady=5)

        # Use grid layout manager and add some margins
        self.neuerKundeButton.grid(row=1, column=0, padx=10, pady=10)
        self.rechnungButton.grid(row=1, column=1, padx=10, pady=10)

        # Configure the columns and row to expand when the window is resized
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=0)
        self.grid_rowconfigure(1, weight=1)

    def open_neuerKundeDialog(self):
        NeuerKundeDialog(self)

    def open_auswahlDialog(self):
        AuswahlDialog(self)


def main():
    VerwaltungListener().mainloop()

if __name__ == "__main__":
    main()